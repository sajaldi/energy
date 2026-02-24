"""
utils.py - Lógica central del sistema de plantillas Word.

Dos funciones principales:
  1. generar_plantilla_en_blanco(model_class) → BytesIO con .docx que contiene
     marcadores {{ campo }} para todos los campos del modelo.

  2. poblar_plantilla(plantilla_word_obj, registro) → BytesIO con .docx
     renderizado con los datos reales del registro.
"""
import io
from django.apps import apps
from django.db import models as django_models


# Campos internos de Django que NO queremos en la plantilla
CAMPOS_EXCLUIDOS = {
    'id', 'creado_en', 'actualizado_en', 'created_at', 'updated_at',
    'contenido_texto', 'hash_archivo', 'estado_extraccion', 'datos_extraidos',
}

# Tipos de campo que manejamos como texto simple
TIPOS_SIMPLES = (
    django_models.CharField,
    django_models.TextField,
    django_models.IntegerField,
    django_models.FloatField,
    django_models.DecimalField,
    django_models.BooleanField,
    django_models.DateField,
    django_models.DateTimeField,
    django_models.EmailField,
    django_models.URLField,
    django_models.SlugField,
)


def _get_campos_modelo(model_class):
    """
    Retorna lista de dicts con info de los campos del modelo.
    Incluye campos simples, FK, M2M y Relaciones Inversas (Related Sets).
    """
    campos = []
    meta = model_class._meta

    # 1. Campos directos (Simples, FK, M2M)
    for field in meta.get_fields():
        name = field.name

        if name in CAMPOS_EXCLUIDOS or name.startswith('_'):
            continue

        # ManyToMany
        if isinstance(field, django_models.ManyToManyField):
            campos.append({
                'nombre': name,
                'etiqueta': getattr(field, 'verbose_name', name),
                'tipo': 'Lista (M2M)',
                'placeholder': f'{{{{ {name}_lista }}}}',
                'codigo': f'{{% for x in {name} %}} {{{{ x.nombre }}}} {{% endfor %}}',
                'descripcion': f'Iterador de {name}. Puedes usar bucles.',
            })
            continue

        # ForeignKey / OneToOne
        if isinstance(field, (django_models.ForeignKey, django_models.OneToOneField)):
            campos.append({
                'nombre': name,
                'etiqueta': str(getattr(field, 'verbose_name', name)),
                'tipo': 'Relación (FK)',
                'placeholder': f'{{{{ {name} }}}}',
                'descripcion': 'Muestra el nombre (__str__) del objeto relacionado.',
            })
            # También permitir acceso a campos del FK: {{ responsable.username }}
            continue

        # Campos simples
        if isinstance(field, TIPOS_SIMPLES):
            campos.append({
                'nombre': name,
                'etiqueta': str(getattr(field, 'verbose_name', name)),
                'tipo': type(field).__name__,
                'placeholder': f'{{{{ {name} }}}}',
                'descripcion': '',
            })

    # 2. Relaciones Inversas (Items relacionados, ej: Revisiones de un Documento)
    for field in meta.get_fields():
        if not field.auto_created or not field.is_relation or field.many_to_many:
            continue
            
        acc_name = field.get_accessor_name()
        if not acc_name or acc_name.endswith('_set') or acc_name in CAMPOS_EXCLUIDOS:
            continue
            
        campos.append({
            'nombre': acc_name,
            'etiqueta': f"Lista de {acc_name.title()}",
            'tipo': 'Tabla/Bucle',
            'placeholder': f'{{% for item in {acc_name} %}} ... {{% endfor %}}',
            'descripcion': f'Permite crear tablas dinámicas con los {acc_name}.',
        })

    # 3. Caso especial: Metadatos Dinámicos (Solo para Documento)
    if model_class.__name__ == 'Documento':
        campos.append({
            'nombre': 'metadatos',
            'etiqueta': 'Metadatos Dinámicos',
            'tipo': 'Variables',
            'placeholder': '{{ m_nombre_metadato }}',
            'descripcion': 'Usa el prefijo "m_" seguido del nombre del metadato configurado.',
        })

    return campos


def generar_plantilla_en_blanco(model_class):
    """
    Genera un archivo Word .docx en blanco con tabla de referencia.
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    doc.add_heading(f'GUÍA DE DISEÑO: {model_class.__name__.upper()}', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph('Variables disponibles para este modelo:')
    p.bold = True

    campos = _get_campos_modelo(model_class)
    tabla = doc.add_table(rows=1, cols=3)
    tabla.style = 'Table Grid'
    hdr = tabla.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = 'Marcador / Bucle', 'Campo', 'Tipo'
    
    for c in hdr:
        for run in c.paragraphs[0].runs: run.bold = True

    for campo in campos:
        row = tabla.add_row().cells
        row[0].text = campo['placeholder']
        row[1].text = campo['etiqueta']
        row[2].text = campo['tipo']

    doc.add_page_break()
    doc.add_heading('--- DISEÑA TU PLANTILLA AQUÍ ---', level=1)
    doc.add_paragraph('Ejemplo de datos básicos:')
    doc.add_paragraph(f"Título: {{{{ titulo }}}}")
    doc.add_paragraph(f"Código: {{{{ codigo }}}}")
    
    doc.add_heading('Ejemplo de Tabla (Iteración):', level=2)
    t_ex = doc.add_table(rows=2, cols=2)
    t_ex.style = 'Table Grid'
    t_ex.rows[0].cells[0].text = '{% for item in revisiones %}' # Estándar Jinja2
    t_ex.rows[0].cells[1].text = 'Rev: {{ item.revision }}'
    t_ex.rows[1].cells[0].text = '{% endfor %}'
    t_ex.rows[1].cells[1].text = '-'

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, campos


def _get_contexto_registro(registro):
    """
    Genera un contexto enriquecido con objetos relacionados para bucles.
    """
    contexto = {}
    model_class = registro.__class__
    
    # 1. Campos base y ForeignKeys
    for field in model_class._meta.fields:
        name = field.name
        val = getattr(registro, name)
        
        if val is None:
            contexto[name] = ''
        elif isinstance(field, (django_models.ForeignKey, django_models.OneToOneField)):
            # Enviamos el objeto completo para que el user pueda hacer {{ responsable.email }}
            contexto[name] = val 
            # Y el string para {{ responsable }}
            contexto[f"{name}_str"] = str(val)
        elif hasattr(val, 'strftime'):
            contexto[name] = val.strftime('%d/%m/%Y')
        else:
            contexto[name] = val

    # 2. ManyToMany (Listas planas)
    for field in model_class._meta.many_to_many:
        name = field.name
        qs = getattr(registro, name).all()
        contexto[f"{name}_lista"] = ", ".join(str(o) for o in qs)
        contexto[name] = qs # Para bucles: {% for x in activos %}

    # 3. Relaciones Inversas (Dinas: tablas)
    # Por ejemplo: 'revisiones', 'comentarios'
    for field in model_class._meta.get_fields():
        if field.auto_created and field.is_relation and not field.many_to_many:
            acc_name = field.get_accessor_name()
            if acc_name and not acc_name.endswith('_set'):
                try:
                    contexto[acc_name] = getattr(registro, acc_name).all()
                except:
                    pass

    # 4. Caso Especial: Metadatos de 'Documento'
    if model_class.__name__ == 'Documento':
        # Inyectamos valores directos: m_nombre_campo
        metadatos = registro.metadatos_valores.select_related('config')
        for mv in metadatos:
            clave = f"m_{mv.config.nombre}"
            contexto[clave] = mv.valor

    return contexto


def poblar_plantilla(plantilla_obj, registro):
    """
    Renderiza la plantilla Word con los datos reales.
    """
    from docxtpl import DocxTemplate
    import tempfile
    import os

    archivo = plantilla_obj.archivo
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        tmp.write(archivo.read())
        tmp_path = tmp.name

    try:
        tpl = DocxTemplate(tmp_path)
        contexto = _get_contexto_registro(registro)
        # Render con soporte para filtros de Jinja2 y bucles de tabla
        tpl.render(contexto)

        buffer = io.BytesIO()
        tpl.save(buffer)
        buffer.seek(0)
        return buffer
    finally:
        os.unlink(tmp_path)

